#!/usr/bin/env python3
"""Converte linhas com pipes em tabelas Word reais no tcc-final.docx."""
from __future__ import annotations
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent
DOCX = ROOT / "docs" / "tcc-final.docx"


def is_pipe_row(text: str) -> bool:
    """Linha com >=2 pipes e celulas nao vazias."""
    if text.count('|') < 2:
        return False
    cells = [c.strip() for c in text.split('|')]
    return len([c for c in cells if c]) >= 2


def find_table_groups(doc) -> list[list[int]]:
    """Encontra grupos de paragrafos consecutivos com pipes."""
    groups = []
    current_group = []
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if is_pipe_row(text):
            current_group.append(i)
        else:
            if len(current_group) >= 2:  # minimo 2 linhas pra virar tabela
                groups.append(current_group)
            current_group = []
    if len(current_group) >= 2:
        groups.append(current_group)
    return groups


def is_number(s: str) -> bool:
    """Detecta se string e numero (pra alinhar a direita)."""
    s = s.strip().replace(',', '.').replace('%', '').replace('ms', '').replace('s', '')
    try:
        float(s)
        return True
    except ValueError:
        return False


def style_cell(cell, text: str, is_header: bool, is_first_col: bool):
    """Aplica formatacao USP/ESALQ a uma celula."""
    cell.text = ''  # limpar
    para = cell.paragraphs[0]

    # Alinhamento: 1a coluna esquerda, cabecalho centralizado, numeros direita, texto justificado
    if is_header:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif is_first_col:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif is_number(text):
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run = para.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Arial'
    # Sem negrito (USP/ESALQ nao permite negrito no conteudo nem no cabecalho)
    run.font.bold = False

    # Espacamento simples (ja e default em tabelas)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)


def create_table_from_rows(doc, rows: list[list[str]]):
    """Cria uma tabela Word a partir das linhas (1a = header)."""
    n_cols = max(len(r) for r in rows)
    # Preencher rows curtas
    rows = [r + [''] * (n_cols - len(r)) for r in rows]

    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'  # bordas simples

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        is_header = (i == 0)
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            style_cell(cell, cell_text, is_header, is_first_col=(j == 0))

    return table


def main():
    doc = Document(str(DOCX))
    groups = find_table_groups(doc)
    print(f"Grupos de tabelas encontrados: {len(groups)}")

    # Processar de TRAS PRA FRENTE (pra nao deslocar indices)
    for group_num, group in enumerate(reversed(groups), 1):
        # Extrair linhas da tabela
        rows = []
        for idx in group:
            text = doc.paragraphs[idx].text.strip()
            cells = [c.strip() for c in text.split('|')]
            # Remover celulas vazias do inicio/fim (markdown pode ter)
            while cells and not cells[0]:
                cells.pop(0)
            while cells and not cells[-1]:
                cells.pop()
            rows.append(cells)

        if not rows:
            continue

        # Criar tabela no FINAL do doc
        new_table = create_table_from_rows(doc, rows)

        # Mover a tabela pra posicao do primeiro paragrafo do grupo
        first_para_el = doc.paragraphs[group[0]]._element
        table_el = new_table._element

        # Remover a tabela do final
        table_el.getparent().remove(table_el)

        # Inserir antes do primeiro paragrafo do grupo
        first_para_el.addprevious(table_el)

        # Remover os paragrafos originais (de tras pra frente)
        for idx in sorted(group, reverse=True):
            para_el = doc.paragraphs[idx]._element
            para_el.getparent().remove(para_el)

        print(f"  Tabela {len(groups) - group_num + 1}: {len(rows)}x{len(rows[0])} celulas")

    doc.save(str(DOCX))
    # Verificar
    doc2 = Document(str(DOCX))
    n_tables = len(doc2.tables)
    print(f"\nSalvo. Total de tabelas no docx: {n_tables}")


if __name__ == '__main__':
    main()
