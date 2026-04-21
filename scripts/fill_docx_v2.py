#!/usr/bin/env python3
"""Preenche tcc-final.docx secao por secao, de TRAS PRA FRENTE pra nao deslocar indices."""
from __future__ import annotations
import re
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent.parent
DOCX_PATH = ROOT / "docs" / "tcc-final.docx"
MD_PATH = ROOT / "docs" / "tcc.md"


def parse_md_sections() -> dict[str, str]:
    text = MD_PATH.read_text()
    text = re.sub(r'<!--.*?-->', '', text, flags=re.DOTALL)
    sections = {}
    current = None
    buf = []
    for line in text.splitlines():
        if line.startswith('## '):
            if current:
                sections[current] = '\n'.join(buf).strip()
            current = line[3:].strip()
            buf = []
        elif current:
            buf.append(line)
    if current:
        sections[current] = '\n'.join(buf).strip()
    return sections


def content_to_paragraphs(content: str) -> list[tuple[str, bool]]:
    """Retorna lista de (texto, is_subtitle). Junta linhas em paragrafos."""
    result = []
    blocks = re.split(r'\n\n+', content)
    for block in blocks:
        block = block.strip()
        if not block or block.startswith('['):
            continue
        # Detectar subtitulo
        if block.startswith('**') and block.endswith('**'):
            title = block.strip('* ')
            result.append((title, True))
            continue
        # Limpar markdown inline
        clean = re.sub(r'\*\*(.+?)\*\*', r'\1', block)
        clean = re.sub(r'\*(.+?)\*', r'\1', clean)
        clean = re.sub(r'`(.+?)`', r'\1', clean)
        # Tabelas markdown → texto com pipes
        if clean.startswith('|'):
            lines = clean.split('\n')
            for line in lines:
                if '---' in line:
                    continue
                cells = [c.strip() for c in line.split('|') if c.strip()]
                result.append((' | '.join(cells), False))
            continue
        # Juntar linhas do mesmo bloco num paragrafo
        clean = ' '.join(clean.split('\n'))
        if clean:
            result.append((clean, False))
    return result


def make_paragraph(ref_element, text: str, bold: bool = False, indent: bool = False):
    """Cria um <w:p> e insere APOS ref_element. Retorna o novo elemento."""
    new_p = ref_element.makeelement(qn('w:p'), {})

    if indent or bold:
        pPr = new_p.makeelement(qn('w:pPr'), {})
        if indent:
            ind = pPr.makeelement(qn('w:ind'), {qn('w:firstLine'): '709'})
            pPr.append(ind)
        new_p.append(pPr)

    run_el = new_p.makeelement(qn('w:r'), {})
    if bold:
        rPr = run_el.makeelement(qn('w:rPr'), {})
        b_el = rPr.makeelement(qn('w:b'), {})
        rPr.append(b_el)
        run_el.append(rPr)

    t_el = run_el.makeelement(qn('w:t'), {})
    t_el.text = text
    t_el.set(qn('xml:space'), 'preserve')
    run_el.append(t_el)
    new_p.append(run_el)

    ref_element.addnext(new_p)
    return new_p


def remove_paragraphs(doc, indices: list[int]):
    """Remove paragrafos pelos indices (de tras pra frente)."""
    for idx in sorted(indices, reverse=True):
        if idx < len(doc.paragraphs):
            el = doc.paragraphs[idx]._element
            el.getparent().remove(el)


def fill_section(doc, title_idx: int, instruction_indices: list[int],
                 paragraphs: list[tuple[str, bool]]):
    """Remove instrucoes e insere conteudo apos o titulo."""
    # Remover instrucoes de tras pra frente
    remove_paragraphs(doc, instruction_indices)

    # Inserir conteudo apos o titulo (que agora pode ter indice diferente)
    # Como removemos paragrafos APOS o titulo, o titulo nao muda de posicao
    title_el = doc.paragraphs[title_idx]._element
    ref = title_el

    for text, is_subtitle in paragraphs:
        ref = make_paragraph(ref, text, bold=is_subtitle, indent=is_subtitle)


def main():
    sections = parse_md_sections()
    doc = Document(str(DOCX_PATH))

    print(f"Secoes no md: {list(sections.keys())}")

    # === PROCESSAR DE TRAS PRA FRENTE (pra nao deslocar indices) ===

    # --- REFERENCIAS (titulo par 73, instrucoes 75-76) ---
    refs_content = sections.get('Referencias', '')
    ref_entries = [(l.strip(), False) for l in refs_content.split('\n')
                   if l.strip() and not l.startswith('[') and not l.startswith('---')]
    # Limpar markdown
    ref_entries = [(re.sub(r'\*(.+?)\*', r'\1', t), False) for t, _ in ref_entries]
    fill_section(doc, 73, [75, 76], ref_entries)
    print(f"  Referencias: {len(ref_entries)} entradas")

    # --- AGRADECIMENTO (titulo par 68, instrucoes 70-71) ---
    fill_section(doc, 68, [70, 71], [])  # Vazio por ora

    # --- CONCLUSAO (titulo par 63, instrucoes 65-66) ---
    conc_paras = content_to_paragraphs(sections.get('Conclusao', ''))
    fill_section(doc, 63, [65, 66], conc_paras)
    # Renomear titulo
    doc.paragraphs[63].clear()
    run = doc.paragraphs[63].add_run('Conclusão')
    run.bold = True
    print(f"  Conclusao: {len(conc_paras)} paragrafos")

    # --- RESULTADOS E DISCUSSAO (titulo par 58, instrucoes 60-61) ---
    rd_paras = content_to_paragraphs(sections.get('Resultados e Discussao', ''))
    fill_section(doc, 58, [56, 60, 61], rd_paras)  # 56 e List Paragraph vazio, 60 instrucao
    print(f"  Resultados: {len(rd_paras)} paragrafos")

    # --- MATERIAL E METODOS (titulo par 49, instrucoes 51-56) ---
    mm_paras = content_to_paragraphs(sections.get('Material e Metodos', ''))
    fill_section(doc, 49, [47, 48, 50, 51, 52], mm_paras)  # 47-52 sao instrucoes
    # Renomear titulo
    doc.paragraphs[49].clear()
    run = doc.paragraphs[49].add_run('Material e Métodos')
    run.bold = True
    print(f"  Material e Metodos: {len(mm_paras)} paragrafos")

    # --- INTRODUCAO (titulo par 43, instrucoes 41, 42, 45-46) ---
    intro_paras = content_to_paragraphs(sections.get('Introducao', ''))
    fill_section(doc, 43, [41, 42, 45, 46], intro_paras)  # 41-42 instrucoes, 45-46 atencao+vazio
    print(f"  Introducao: {len(intro_paras)} paragrafos")

    # --- ABSTRACT (titulo par 38, conteudo par 40, keywords 41) ---
    abstract_paras = content_to_paragraphs(sections.get('Abstract', ''))
    doc.paragraphs[38].clear()
    run = doc.paragraphs[38].add_run('Abstract')
    run.bold = True
    # Substituir pontos por texto do abstract
    if abstract_paras:
        doc.paragraphs[40].clear()
        doc.paragraphs[40].add_run(abstract_paras[0][0])
    # Keywords
    doc.paragraphs[37].clear()  # "Keywords ou Palabras Clave"
    run = doc.paragraphs[37].add_run('Keywords: microservices; serverless; Domain-Driven Design; static analysis; productivity')
    run.bold = True
    print(f"  Abstract: {len(abstract_paras)} paragrafos")

    # --- RESUMO (titulo par 30, conteudo par 32, palavras-chave 33) ---
    resumo_paras = content_to_paragraphs(sections.get('Resumo', ''))
    doc.paragraphs[30].clear()
    run = doc.paragraphs[30].add_run('Resumo')
    run.bold = True
    if resumo_paras:
        doc.paragraphs[32].clear()
        doc.paragraphs[32].add_run(resumo_paras[0][0])
    doc.paragraphs[33].clear()
    run = doc.paragraphs[33].add_run('Palavras-chave: microsserviços; "serverless"; "Domain-Driven Design"; análise estática; produtividade')
    run.bold = True
    # Remover "Atencao" (par 34)
    doc.paragraphs[34]._element.getparent().remove(doc.paragraphs[34]._element)
    print(f"  Resumo: preenchido")

    # --- TITULO (par 0 e 28) ---
    titulo = 'Migração de monolito para microsserviços serverless com Domain-Driven Design'
    doc.paragraphs[0].clear()
    run = doc.paragraphs[0].add_run(titulo)
    run.bold = True
    doc.paragraphs[28].clear()
    run = doc.paragraphs[28].add_run(titulo)
    run.bold = True

    # Titulo ingles (par 36)
    doc.paragraphs[36].clear()
    run = doc.paragraphs[36].add_run('Migration from monolith to serverless microservices with Domain-Driven Design')
    run.bold = True

    # --- AUTORES (par 3, 5, 6, 7) ---
    doc.paragraphs[3].clear()
    doc.paragraphs[3].add_run('Gabriel Figueiredo dos Santos¹*; Ugo Henrique Pereira da Silva²')
    doc.paragraphs[5].clear()
    doc.paragraphs[5].add_run('¹ Engenheiro de Software – São Paulo, SP, Brasil')
    doc.paragraphs[6].clear()
    doc.paragraphs[6].add_run('² Universidade de São Paulo, ESALQ – Piracicaba, SP, Brasil')
    doc.paragraphs[7].clear()
    doc.paragraphs[7].add_run('*autor correspondente: gabrielfgsantos95@gmail.com')

    # === LIMPAR SOBRAS ===
    # Remover todas as instrucoes restantes do template
    to_remove = []
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if 'Atenção: antes de enviar' in text:
            to_remove.append(i)
        elif 'Metodologia ou Material e Métodos' in text and 'título da seção' in text.lower():
            to_remove.append(i)
    remove_paragraphs(doc, to_remove)

    # === SALVAR ===
    doc.save(str(DOCX_PATH))
    total = sum(1 for p in doc.paragraphs if p.text.strip())
    print(f"\nSalvo em {DOCX_PATH}")
    print(f"Total paragrafos com texto: {total}")


if __name__ == '__main__':
    main()
