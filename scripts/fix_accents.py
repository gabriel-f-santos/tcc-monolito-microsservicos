#!/usr/bin/env python3
"""Corrige acentuacao em portugues no tcc-final.docx e tcc.md."""
from __future__ import annotations
import re
from pathlib import Path
from docx import Document

ROOT = Path(__file__).resolve().parent.parent
DOCX = ROOT / "docs" / "tcc-final.docx"
MD = ROOT / "docs" / "tcc.md"


# Dicionario sem_acento → com_acento (ordem importa pra palavras mais especificas primeiro)
ACCENTS = {
    # substantivos -acao → -ação
    "migracao": "migração", "migracoes": "migrações",
    "implementacao": "implementação", "implementacoes": "implementações",
    "integracao": "integração", "integracoes": "integrações",
    "comunicacao": "comunicação", "comunicacoes": "comunicações",
    "aplicacao": "aplicação", "aplicacoes": "aplicações",
    "execucao": "execução", "execucoes": "execuções",
    "configuracao": "configuração", "configuracoes": "configurações",
    "apresentacao": "apresentação", "apresentacoes": "apresentações",
    "documentacao": "documentação",
    "validacao": "validação", "validacoes": "validações",
    "transformacao": "transformação",
    "orientacao": "orientação",
    "utilizacao": "utilização",
    "criacao": "criação", "criacoes": "criações",
    "definicao": "definição", "definicoes": "definições",
    "classificacao": "classificação",
    "modificacao": "modificação", "modificacoes": "modificações",
    "manutencao": "manutenção",
    "producao": "produção",
    "observacao": "observação", "observacoes": "observações",
    "avaliacao": "avaliação", "avaliacoes": "avaliações",
    "operacao": "operação", "operacoes": "operações",
    "publicacao": "publicação", "publicacoes": "publicações",
    "instabilidade": "instabilidade",  # OK
    "duplicacao": "duplicação",
    "reutilizacao": "reutilização",
    "decomposicao": "decomposição",
    "especializacao": "especialização",
    "otimizacao": "otimização",
    "direcao": "direção",
    "funcao": "função", "funcoes": "funções",
    "solucao": "solução", "solucoes": "soluções",
    "projecao": "projeção", "projecoes": "projeções",
    "decisao": "decisão", "decisoes": "decisões",
    "questao": "questão", "questoes": "questões",
    "conclusao": "conclusão", "conclusoes": "conclusões",
    "excecao": "exceção", "excecoes": "exceções",
    "inclusao": "inclusão",
    "revisao": "revisão", "revisoes": "revisões",
    "divisao": "divisão", "divisoes": "divisões",
    "pressao": "pressão",
    "tensao": "tensão",
    "extensao": "extensão", "extensoes": "extensões",
    "compreensao": "compreensão",
    "expressao": "expressão",
    "conexao": "conexão", "conexoes": "conexões",
    "discussao": "discussão",

    # -encia → -ência
    "referencia": "referência", "referencias": "referências",
    "tendencia": "tendência", "tendencias": "tendências",
    "consistencia": "consistência",
    "dependencia": "dependência", "dependencias": "dependências",
    "frequencia": "frequência",
    "sequencia": "sequência",
    "convergencia": "convergência",
    "presenca": "presença",
    "ausencia": "ausência",
    "existencia": "existência",
    "influencia": "influência",
    "experiencia": "experiência", "experiencias": "experiências",
    "ciencia": "ciência",
    "eficiencia": "eficiência",
    "evidencia": "evidência", "evidencias": "evidências",
    "ocorrencia": "ocorrência", "ocorrencias": "ocorrências",
    "diferencia": "diferença",  # typo? fix to diferenca
    "diferenca": "diferença", "diferencas": "diferenças",
    "latencia": "latência", "latencias": "latências",
    "emergencia": "emergência",
    "potencia": "potência",
    "resistencia": "resistência",

    # palavras comuns
    "metrica": "métrica", "metricas": "métricas",
    "tecnica": "técnica", "tecnicas": "técnicas",
    "tecnico": "técnico", "tecnicos": "técnicos",
    "criterio": "critério", "criterios": "critérios",
    "especifico": "específico", "especifica": "específica",
    "especificos": "específicos", "especificas": "específicas",
    "pratica": "prática", "praticas": "práticas",
    "pratico": "prático", "praticos": "práticos",
    "teorico": "teórico", "teorica": "teórica",
    "teoricos": "teóricos", "teoricas": "teóricas",
    "analise": "análise", "analises": "análises",
    "sintese": "síntese",
    "hipotese": "hipótese", "hipoteses": "hipóteses",
    "pagina": "página", "paginas": "páginas",
    "codigo": "código", "codigos": "códigos",
    "dominio": "domínio", "dominios": "domínios",
    "indice": "índice", "indices": "índices",
    "ciclomatica": "ciclomática",
    "monolitico": "monolítico", "monolitica": "monolítica",
    "monoliticos": "monolíticos", "monoliticas": "monolíticas",
    "monoliticamente": "monoliticamente",
    "servico": "serviço", "servicos": "serviços",
    "microsservico": "microsserviço", "microsservicos": "microsserviços",
    "microsservices": "microsserviços",
    "servidor": "servidor", "servidores": "servidores",  # OK
    "memoria": "memória", "memorias": "memórias",
    "historia": "história", "historias": "histórias",
    "historico": "histórico", "historica": "histórica",
    "historicos": "históricos", "historicas": "históricas",
    "cenario": "cenário", "cenarios": "cenários",
    "conteudo": "conteúdo", "conteudos": "conteúdos",
    "inicio": "início",
    "proprio": "próprio", "propria": "própria",
    "proprios": "próprios", "proprias": "próprias",
    "relatorio": "relatório", "relatorios": "relatórios",
    "repositorio": "repositório", "repositorios": "repositórios",
    "diretorio": "diretório", "diretorios": "diretórios",
    "numero": "número", "numeros": "números",
    "modulo": "módulo", "modulos": "módulos",
    "formula": "fórmula", "formulas": "fórmulas",
    "unico": "único", "unica": "única",
    "unicos": "únicos", "unicas": "únicas",
    "ultimo": "último", "ultima": "última",
    "ultimos": "últimos", "ultimas": "últimas",
    "minimo": "mínimo", "minimos": "mínimos",
    "maximo": "máximo", "maximos": "máximos",
    "proximo": "próximo", "proxima": "próxima",
    "proximos": "próximos", "proximas": "próximas",
    "basico": "básico", "basica": "básica",
    "basicos": "básicos", "basicas": "básicas",
    "classico": "clássico", "classica": "clássica",
    "classicos": "clássicos", "classicas": "clássicas",
    "tipico": "típico", "tipica": "típica",
    "tipicos": "típicos", "tipicas": "típicas",
    "logico": "lógico", "logica": "lógica",
    "logicos": "lógicos", "logicas": "lógicas",
    "publico": "público", "publica": "pública",
    "publicos": "públicos", "publicas": "públicas",
    "academico": "acadêmico", "academica": "acadêmica",
    "academicos": "acadêmicos", "academicas": "acadêmicas",
    "sistematico": "sistemático", "sistematica": "sistemática",
    "automatico": "automático", "automatica": "automática",
    "automaticos": "automáticos", "automaticas": "automáticas",
    "automatizada": "automatizada", "automatizado": "automatizado",
    "estatistica": "estatística", "estatico": "estático", "estatica": "estática",
    "domestico": "doméstico", "domestica": "doméstica",
    "organico": "orgânico", "organica": "orgânica",
    "dinamico": "dinâmico", "dinamica": "dinâmica",
    "economico": "econômico", "economica": "econômica",
    "sinonimo": "sinônimo", "anonimo": "anônimo",
    "facil": "fácil", "faceis": "fáceis",
    "dificil": "difícil", "dificeis": "difíceis",
    "util": "útil", "uteis": "úteis",
    "possivel": "possível", "possiveis": "possíveis",
    "impossivel": "impossível",
    "viavel": "viável", "viaveis": "viáveis",
    "agil": "ágil", "ageis": "ágeis",
    "movel": "móvel", "moveis": "móveis",
    "nivel": "nível", "niveis": "níveis",
    "disponivel": "disponível", "disponiveis": "disponíveis",
    "variavel": "variável", "variaveis": "variáveis",
    "responsavel": "responsável", "responsaveis": "responsáveis",
    "aplicavel": "aplicável",
    "generico": "genérico", "generica": "genérica",
    "superficie": "superfície",
    "superficial": "superficial", "superficiais": "superficiais",
    "necessario": "necessário", "necessaria": "necessária",
    "necessarios": "necessários", "necessarias": "necessárias",
    "diario": "diário", "diaria": "diária",
    "ordinario": "ordinário", "ordinaria": "ordinária",
    "obrigatorio": "obrigatório", "obrigatoria": "obrigatória",

    # outras palavras comuns
    "porem": "porém",
    "alem": "além",
    "apos": "após",
    "ate": "até",
    "possui": "possui",  # OK
    "so": "só",
    "nao": "não",
    "ja": "já",
    "tambem": "também",
    "e os": "é os",  # tricky
    "atraves": "através",
    "varias": "várias", "varios": "vários",
    "proprio": "próprio",
    "caracteristica": "característica", "caracteristicas": "características",
    "produtos": "produtos",  # OK
    "estoque": "estoque",  # OK
    "categoria": "categoria",  # OK
    "categorias": "categorias",  # OK
    "registrar": "registrar",  # OK
    "exceto": "exceto",  # OK
    "pre": "pré",
    "pos": "pós",
    "tres": "três",
    "decada": "década",
    "maioria": "maioria",  # OK
    "similaridade": "similaridade",  # OK
    "ciencias": "ciências",
    "experimentacao": "experimentação",
    "desempenho": "desempenho",  # OK
    "rapido": "rápido", "rapida": "rápida",
    "rapidos": "rápidos", "rapidas": "rápidas",
    "ocorreu": "ocorreu",  # OK
    "autor": "autor",  # OK
    "autores": "autores",  # OK
    "autonomo": "autônomo", "autonoma": "autônoma",
    "autonomos": "autônomos", "autonomas": "autônomas",
    "explicito": "explícito", "explicita": "explícita",
    "explicitos": "explícitos", "explicitas": "explícitas",
    "implicito": "implícito",
    "familia": "família", "familias": "famílias",
    "familiar": "familiar",  # OK
    "distribuido": "distribuído", "distribuida": "distribuída",
    "distribuidos": "distribuídos", "distribuidas": "distribuídas",
    "construido": "construído", "construida": "construída",
    "contribuiu": "contribuiu",  # OK
    "volatil": "volátil", "volateis": "voláteis",
    "util": "útil",
    "estavel": "estável", "estaveis": "estáveis",
    "estabil": "estável",
    "estabilidade": "estabilidade",  # OK
    "perfil": "perfil",  # OK
    "artigo": "artigo",  # OK
    "codigos": "códigos",
    "bibliografica": "bibliográfica", "bibliograficas": "bibliográficas",
    "cronometrado": "cronometrado",  # OK
    "cronometragem": "cronometragem",  # OK
    "criterio": "critério",
    "criterios": "critérios",
    "intencional": "intencional",  # OK
    "propriedades": "propriedades",  # OK
    "automatizados": "automatizados",  # OK
    "automatizado": "automatizado",  # OK
    "condiciono": "condicionou",  # typo fix
    "e uma": "é uma", "e o": "é o", "e a": "é a", "e um": "é um",

    # segunda rodada - palavras que escaparam
    "inteligencia": "inteligência",
    "catalogo": "catálogo",  # mas projeto usa "Catalogo" como nome — cuidado!
    "media": "média", "medias": "médias",
    "dispersao": "dispersão",
    "iteracao": "iteração", "iteracoes": "iterações",
    "mensuravel": "mensurável", "mensuraveis": "mensuráveis",
    "evolucao": "evolução",
    "introducao": "introdução",
    "automatizacao": "automatização",
    "verificacao": "verificação",
    "padronizacao": "padronização",
    "organizacao": "organização", "organizacoes": "organizações",
    "justificativa": "justificativa",  # OK
    "adequada": "adequada",  # OK
    "adequadamente": "adequadamente",  # OK
    "adequado": "adequado",  # OK
    "percentil": "percentil",  # OK
    "percentis": "percentis",  # OK
    "percepcao": "percepção",
    "resolucao": "resolução",
    "persistencia": "persistência",
    "assistencia": "assistência",
    "instancia": "instância", "instancias": "instâncias",
    "distancia": "distância", "distancias": "distâncias",
    "circunstancia": "circunstância",
    "importancia": "importância",
    "relevancia": "relevância",
    "abrangencia": "abrangência",
    "negocio": "negócio", "negocios": "negócios",
    "proposito": "propósito",
    "proposta": "proposta", "propostas": "propostas",  # OK
    "proposto": "proposto", "propostos": "propostos",  # OK
    "teste": "teste", "testes": "testes",  # OK
    "estrutura": "estrutura",  # OK
    "estruturada": "estruturada",  # OK
    "estrutural": "estrutural",  # OK
    "estruturais": "estruturais",  # OK
    "conhecimento": "conhecimento",  # OK
    "conhecimentos": "conhecimentos",  # OK
    "desempenho": "desempenho",  # OK
    "autenticacao": "autenticação",
    "autenticado": "autenticado",  # OK
    "acoplamento": "acoplamento",  # OK
    "aferente": "aferente",  # OK
    "eferente": "eferente",  # OK
    "atomica": "atômica", "atomicos": "atômicos",
    "aleatorio": "aleatório", "aleatoria": "aleatória",
    "automatizada": "automatizada",  # OK
    "opcoes": "opções",
    "versoes": "versões",
    "razoes": "razões",
    "razao": "razão",
    "escolha": "escolha",  # OK
    "escolhas": "escolhas",  # OK
    "escolhido": "escolhido",  # OK
    "metade": "metade",  # OK
    "metodo": "método", "metodos": "métodos",
    "metodologia": "metodologia",  # OK
    "metodologias": "metodologias",  # OK
    "metodologico": "metodológico", "metodologica": "metodológica",
    "processo": "processo",  # OK
    "processos": "processos",  # OK
    "processamento": "processamento",  # OK
    "nivel": "nível",
    "niveis": "níveis",
    "util": "útil",
    "uteis": "úteis",
    "eventos": "eventos",  # OK
    "implantacao": "implantação",
    "autonomia": "autonomia",  # OK
    "porta": "porta",  # OK (pode ser conflito se citar "porta")
    "versao": "versão", "versoes": "versões",
    "conexao": "conexão",
    "especial": "especial",  # OK
    "especialmente": "especialmente",  # OK
    "especializado": "especializado",
    "conhecimento": "conhecimento",  # OK
    "ubiqua": "ubíqua",
    "agregacao": "agregação",
    "agregacoes": "agregações",
    "composto": "composto",  # OK
    "composta": "composta",  # OK
    "composicoes": "composições",
    "composicao": "composição",
    "mantem": "mantém",
    "conhece": "conhece",  # OK
    "conforme": "conforme",  # OK
    "implica": "implica",  # OK
    "implicacoes": "implicações",
    "validacao": "validação",
    "persistencia": "persistência",
    "semantica": "semântica",
    "delegacao": "delegação",
    "rotulos": "rótulos",
    "rotulo": "rótulo",
    "topologia": "topologia",  # OK
    "topologias": "topologias",  # OK
    "topico": "tópico", "topicos": "tópicos",
    "subtopico": "subtópico", "subtopicos": "subtópicos",
    "eletronica": "eletrônica",
    "eletronico": "eletrônico",
    "tabela": "tabela",  # OK
    "tabelas": "tabelas",  # OK
    "acentos": "acentos",  # OK
    "acentuacao": "acentuação",
    "estetica": "estética",
    "estaticas": "estáticas",  # OK
    "estatico": "estático",
    "estatica": "estática",
    "cristalico": "cristálico",
    "sumario": "sumário",
    "sumarios": "sumários",
    "objetivo": "objetivo",  # OK
    "objetivos": "objetivos",  # OK
    "categoria": "categoria",  # OK
    "categorias": "categorias",  # OK
    "cripto": "cripto",  # OK
    "criptografia": "criptografia",  # OK
    "abstratos": "abstratos",  # OK
    "abstrato": "abstrato",  # OK
    "abstracao": "abstração",
    "abstracoes": "abstrações",
    "projeto": "projeto",  # OK
    "projetos": "projetos",  # OK
    "projetar": "projetar",  # OK
    "projetada": "projetada",  # OK
    "projetado": "projetado",  # OK
    "sessao": "sessão", "sessoes": "sessões",
    "implicacao": "implicação",
    "implicacoes": "implicações",
    "repeticao": "repetição",
    "repeticoes": "repetições",
    "funcional": "funcional",  # OK
    "funcionais": "funcionais",  # OK
    "funcionalidade": "funcionalidade",  # OK
    "funcionalidades": "funcionalidades",  # OK
    "fronteira": "fronteira",  # OK
    "fronteiras": "fronteiras",  # OK
    "horizonte": "horizonte",  # OK
    "horizontes": "horizontes",  # OK
    "horizontais": "horizontais",  # OK
    "horizontal": "horizontal",  # OK
    "vertical": "vertical",  # OK
    "verticais": "verticais",  # OK
    "condicao": "condição",
    "condicoes": "condições",
    "estocastica": "estocástica",
    "replicacao": "replicação",
    "prolongada": "prolongada",  # OK
    "realizar": "realizar",  # OK
    "realizada": "realizada",  # OK
    "realizado": "realizado",  # OK
    "sufixo": "sufixo",  # OK
    "prefixo": "prefixo",  # OK
    "exterior": "exterior",  # OK
    "interior": "interior",  # OK
    "posicao": "posição",
    "posicoes": "posições",
    "previsibilidade": "previsibilidade",  # OK
    "previsivel": "previsível", "previsiveis": "previsíveis",
    "infraestrutura": "infraestrutura",  # OK
    "infraestruturas": "infraestruturas",  # OK

    # terceira rodada
    "comparacao": "comparação", "comparacoes": "comparações",
    "computacao": "computação",
    "movimentacao": "movimentação", "movimentacoes": "movimentações",
    "simulacao": "simulação",
    "adequacao": "adequação",
    "alteracao": "alteração", "alteracoes": "alterações",
    "concentracao": "concentração",
    "degradacao": "degradação",
    "derivacao": "derivação",
    "duracao": "duração",
    "emulacao": "emulação",
    "especificacao": "especificação", "especificacoes": "especificações",
    "formatacao": "formatação",
    "limitacao": "limitação", "limitacoes": "limitações",
    "marcacao": "marcação", "marcacoes": "marcações",
    "ramificacao": "ramificação", "ramificacoes": "ramificações",
    "concorrencia": "concorrência",
    "consequencia": "consequência", "consequencias": "consequências",
    "gerencia": "gerência",
    "resiliencia": "resiliência",
    "preferencia": "preferência", "preferencias": "preferências",
    "comparavel": "comparável", "comparaveis": "comparáveis",
    "implantavel": "implantável", "implantaveis": "implantáveis",
    "imutavel": "imutável", "imutaveis": "imutáveis",
    "monitoravel": "monitorável", "monitoraveis": "monitoráveis",
    "executavel": "executável",
    "favoravel": "favorável",
    "tratavel": "tratável",
    "intratavel": "intratável",
    "justificavel": "justificável",
    "mensuravel": "mensurável",
    "controlavel": "controlável",
    "observavel": "observável",
    "especializavel": "especializável",
    "desativavel": "desativável",

    # quarta rodada
    "formulacao": "formulação",
    "generalizacao": "generalização",
    "geracao": "geração",
    "importacao": "importação",
    "inicializacao": "inicialização",
    "interpretacao": "interpretação",
    "mitigacao": "mitigação",
    "notacao": "notação",
    "separacao": "separação", "separacoes": "separações",
    "variacao": "variação", "variacoes": "variações",
    "iteracao": "iteração", "iteracoes": "iterações",
    "invocacao": "invocação", "invocacoes": "invocações",
    "instalacao": "instalação",
    "navegacao": "navegação",
    "renovacao": "renovação",
    "restricao": "restrição", "restricoes": "restrições",
    "criticacao": "criticação",
    "reducao": "redução",
    "exposicao": "exposição",
    "inspecao": "inspeção", "inspecoes": "inspeções",
    "selecao": "seleção", "selecoes": "seleções",
    "correcao": "correção", "correcoes": "correções",
    "contribuicao": "contribuição", "contribuicoes": "contribuições",
    "atualizacao": "atualização", "atualizacoes": "atualizações",
    "simplificacao": "simplificação",
    "transmissao": "transmissão",
    "colisao": "colisão",
    "caracterizacao": "caracterização",
    "especializacao": "especialização",
    "emissao": "emissão",
    "permissao": "permissão", "permissoes": "permissões",
    "impressao": "impressão",
    "concessao": "concessão",
    "elaboracao": "elaboração",
    "recuperacao": "recuperação",
    "falacia": "falácia",
    "auditoria": "auditoria",  # OK

    # quinta rodada
    "agnostica": "agnóstica", "agnostico": "agnóstico",
    "identica": "idêntica", "identico": "idêntico",
    "identicas": "idênticas", "identicos": "idênticos",
    "sintetica": "sintética", "sintetico": "sintético",
    "criptografico": "criptográfico", "criptografica": "criptográfica",
    "grafico": "gráfico", "grafica": "gráfica",
    "graficos": "gráficos", "graficas": "gráficas",
    "empirico": "empírico", "empirica": "empírica",
    "cientifico": "científico", "cientifica": "científica",
    "periodico": "periódico", "periodicos": "periódicos",
    "metalurgico": "metalúrgico",
    "cenarios": "cenários",
    "tamanho": "tamanho",  # OK
    "pequeno": "pequeno",  # OK
    "segundo": "segundo",  # OK
    "contagem": "contagem",  # OK
}

# Regex para palavras inteiras (respeitando capitalização)
def build_regex():
    # Ordenar por tamanho decrescente pra matches mais longos primeiro
    sorted_keys = sorted(ACCENTS.keys(), key=len, reverse=True)
    pattern = r'\b(' + '|'.join(re.escape(k) for k in sorted_keys) + r')\b'
    return re.compile(pattern, re.IGNORECASE)


PATTERN = build_regex()


def preserve_case(original: str, replacement: str) -> str:
    """Preserva capitalizacao do original."""
    if original.isupper():
        return replacement.upper()
    elif original[0].isupper():
        return replacement[0].upper() + replacement[1:]
    return replacement


def fix_text(text: str) -> str:
    """Aplica substituicoes ao texto."""
    def repl(m):
        orig = m.group(0)
        low = orig.lower()
        rep = ACCENTS.get(low, orig)
        return preserve_case(orig, rep)
    return PATTERN.sub(repl, text)


def fix_docx():
    doc = Document(str(DOCX))
    total_changes = 0

    # Paragrafos normais
    for p in doc.paragraphs:
        for run in p.runs:
            new_text = fix_text(run.text)
            if new_text != run.text:
                total_changes += 1
                run.text = new_text

    # Celulas de tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        new_text = fix_text(run.text)
                        if new_text != run.text:
                            total_changes += 1
                            run.text = new_text

    doc.save(str(DOCX))
    print(f"DOCX: {total_changes} runs corrigidos")


def fix_md():
    text = MD.read_text()
    new_text = fix_text(text)
    changes = sum(1 for a, b in zip(text, new_text) if a != b)
    MD.write_text(new_text)
    print(f"MD: {changes} caracteres modificados")


if __name__ == '__main__':
    fix_md()
    fix_docx()
    print("\nCorrecao de acentos aplicada. Revisar visualmente casos especificos.")
