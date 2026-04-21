#!/usr/bin/env python3
"""Preenche tcc-final.docx com conteudo do tcc.md, preservando estilos do template."""
from __future__ import annotations
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parent.parent
DOCX_PATH = ROOT / "docs" / "tcc-final.docx"
MD_PATH = ROOT / "docs" / "tcc.md"
OUT_PATH = ROOT / "docs" / "tcc-final.docx"

# --- Parse tcc.md ---

def parse_md() -> dict[str, str]:
    """Extrai secoes do tcc.md. Retorna {nome_secao: conteudo}."""
    text = MD_PATH.read_text()
    # Remove comentarios HTML
    text = re.sub(r'<!--.*?-->', '', text, flags=re.DOTALL)

    sections = {}
    current = None
    buf = []

    for line in text.splitlines():
        if line.startswith('## '):
            if current:
                sections[current] = '\n'.join(buf).strip()
            current = line[3:].strip()
            buf = []
        elif current:
            buf.append(line)
    if current:
        sections[current] = '\n'.join(buf).strip()

    return sections


def extract_section_paragraphs(content: str) -> list[dict]:
    """Converte conteudo markdown em lista de paragrafos com metadata."""
    paras = []
    lines = content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]

        # Pular linhas de tabela markdown (|---|)
        if line.strip().startswith('|') and '---' in line:
            i += 1
            continue

        # Tabela markdown - converter pra texto simples por ora
        if line.strip().startswith('|'):
            # Acumular toda a tabela
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                if '---' not in lines[i]:
                    table_lines.append(lines[i])
                i += 1
            # Marcar como tabela
            paras.append({'type': 'table', 'lines': table_lines})
            continue

        # Subtitulo (negrito markdown)
        if line.strip().startswith('**') and line.strip().endswith('**'):
            title = line.strip().strip('*').strip()
            paras.append({'type': 'subtitle', 'text': title})
            i += 1
            continue

        # Paragrafo normal (pode ter multiplas linhas ate linha vazia)
        if line.strip():
            para_text = line.strip()
            # Nao juntar linhas — cada linha markdown e um paragrafo
            # (o md ja tem paragrafos separados por linhas vazias)
            paras.append({'type': 'normal', 'text': para_text})

        i += 1

    return paras


# --- Manipular docx ---

def clear_range(doc: Document, start_idx: int, end_idx: int):
    """Remove paragrafos entre start_idx e end_idx (inclusive)."""
    # Marcar pra remover (de tras pra frente)
    for i in range(end_idx, start_idx - 1, -1):
        if i < len(doc.paragraphs):
            p = doc.paragraphs[i]._element
            p.getparent().remove(p)


def insert_paragraph_after(doc: Document, ref_idx: int, text: str,
                           style: str = 'Normal', bold: bool = False,
                           first_line_indent: float | None = None) -> int:
    """Insere paragrafo apos o paragrafo ref_idx. Retorna novo indice."""
    ref_para = doc.paragraphs[ref_idx]
    new_para = ref_para._element.addnext(
        ref_para._element.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p', {})
    )
    # Recarregar documento pra pegar o novo paragrafo
    # (python-docx nao atualiza a lista automaticamente)
    return ref_idx + 1


def fill_section(doc: Document, section_title_idx: int, content: str,
                 instruction_start: int, instruction_end: int):
    """Preenche uma secao do docx com conteudo do md."""
    paras = extract_section_paragraphs(content)

    # Remover instrucoes do template (de tras pra frente)
    for i in range(instruction_end, instruction_start - 1, -1):
        if i < len(doc.paragraphs):
            el = doc.paragraphs[i]._element
            el.getparent().remove(el)

    # Inserir conteudo apos o titulo da secao
    ref = doc.paragraphs[section_title_idx]._element

    for para in reversed(paras):
        from docx.oxml.ns import qn
        new_p = ref.makeelement(qn('w:p'), {})

        if para['type'] == 'subtitle':
            # Paragrafo com subtitulo em negrito
            pPr = new_p.makeelement(qn('w:pPr'), {})
            # Recuo 1.25cm na primeira linha
            ind = pPr.makeelement(qn('w:ind'), {qn('w:firstLine'): '709'})  # 709 twips = 1.25cm
            pPr.append(ind)
            new_p.append(pPr)

            run = new_p.makeelement(qn('w:r'), {})
            rPr = run.makeelement(qn('w:rPr'), {})
            bold_el = rPr.makeelement(qn('w:b'), {})
            rPr.append(bold_el)
            run.append(rPr)

            t = run.makeelement(qn('w:t'), {})
            t.text = para['text']
            run.append(t)
            new_p.append(run)

        elif para['type'] == 'table':
            # Tabelas markdown — inserir como texto formatado por ora
            # (usuario pode converter pra tabela Word depois)
            for tline in reversed(para['lines']):
                cells = [c.strip() for c in tline.split('|') if c.strip()]
                row_text = ' | '.join(cells)
                tp = ref.makeelement(qn('w:p'), {})
                run = tp.makeelement(qn('w:r'), {})
                t = run.makeelement(qn('w:t'), {})
                t.text = row_text
                t.set(qn('xml:space'), 'preserve')
                run.append(t)
                tp.append(run)
                ref.addnext(tp)
            continue

        else:
            # Paragrafo normal
            run = new_p.makeelement(qn('w:r'), {})
            t = run.makeelement(qn('w:t'), {})
            # Limpar markdown inline
            clean = para['text']
            clean = re.sub(r'\*\*(.+?)\*\*', r'\1', clean)  # bold
            clean = re.sub(r'\*(.+?)\*', r'\1', clean)  # italic
            clean = re.sub(r'`(.+?)`', r'\1', clean)  # code
            t.text = clean
            t.set(qn('xml:space'), 'preserve')
            run.append(t)
            new_p.append(run)

        ref.addnext(new_p)


def main():
    sections = parse_md()
    doc = Document(str(DOCX_PATH))

    print(f"Secoes encontradas no md: {list(sections.keys())}")

    # --- Titulo (paragrafo 0 e 24) ---
    novo_titulo = "Migração de monolito para microsserviços serverless com Domain-Driven Design"
    doc.paragraphs[0].clear()
    run = doc.paragraphs[0].add_run(novo_titulo)
    run.bold = True

    if len(doc.paragraphs) > 24:
        doc.paragraphs[24].clear()
        run = doc.paragraphs[24].add_run(novo_titulo)
        run.bold = True

    # --- Resumo (paragrafo 28 = instrucoes → substituir) ---
    if 'Resumo' in sections:
        resumo_text = sections['Resumo']
        # Extrair so o paragrafo principal (sem placeholder, sem palavras-chave)
        lines = [l for l in resumo_text.split('\n') if l.strip() and not l.startswith('**Palavras') and not l.startswith('[')]
        if lines:
            doc.paragraphs[28].clear()
            run = doc.paragraphs[28].add_run(lines[0])
            run.bold = False

        # Palavras-chave (paragrafo 29)
        doc.paragraphs[29].clear()
        run = doc.paragraphs[29].add_run('Palavras-chave: microsserviços; "serverless"; "Domain-Driven Design"; análise estática; produtividade')
        run.bold = True

    # Remover "Atencao" (paragrafo 30)
    doc.paragraphs[30]._element.getparent().remove(doc.paragraphs[30]._element)

    # Recarregar indices apos remocao
    doc = Document(str(DOCX_PATH))  # Nao funciona assim, preciso salvar e recarregar

    # Abordagem diferente: marcar tudo e salvar de uma vez
    # Vou usar abordagem mais segura - preencher sem remover, so substituir texto

    doc = Document(str(DOCX_PATH))

    # === TITULO ===
    doc.paragraphs[0].clear()
    run = doc.paragraphs[0].add_run(novo_titulo)
    run.bold = True

    doc.paragraphs[24].clear()
    run = doc.paragraphs[24].add_run(novo_titulo)
    run.bold = True

    # === RESUMO (par 28) ===
    resumo_lines = [l.strip() for l in sections.get('Resumo', '').split('\n')
                    if l.strip() and not l.startswith('**') and not l.startswith('[') and not l.startswith('<!--')]
    if resumo_lines:
        doc.paragraphs[28].clear()
        doc.paragraphs[28].add_run(' '.join(resumo_lines))

    # === PALAVRAS-CHAVE (par 29) ===
    doc.paragraphs[29].clear()
    run = doc.paragraphs[29].add_run('Palavras-chave: microsserviços; "serverless"; "Domain-Driven Design"; análise estática; produtividade')
    run.bold = True

    # === ATENCAO resumo (par 30) → limpar ===
    doc.paragraphs[30].clear()

    # === TITULO INGLES (par 32) ===
    doc.paragraphs[32].clear()
    run = doc.paragraphs[32].add_run('Migration from monolith to serverless microservices with Domain-Driven Design')
    run.bold = True

    # === ABSTRACT (par 34 = titulo, par 36 = conteudo) ===
    doc.paragraphs[34].clear()
    run = doc.paragraphs[34].add_run('Abstract')
    run.bold = True

    abstract_lines = [l.strip() for l in sections.get('Abstract', '').split('\n')
                      if l.strip() and not l.startswith('**') and not l.startswith('[') and not l.startswith('<!--')]
    if abstract_lines:
        doc.paragraphs[36].clear()
        doc.paragraphs[36].add_run(' '.join(abstract_lines))

    # === KEYWORDS (par 37) ===
    doc.paragraphs[37].clear()
    run = doc.paragraphs[37].add_run('Keywords: microservices; serverless; Domain-Driven Design; static analysis; productivity')
    run.bold = True

    # === INTRODUCAO (par 39 = titulo MANTER, par 41-43 = instrucoes → substituir) ===
    intro_content = sections.get('Introducao', '')
    intro_paras = [l.strip() for l in intro_content.split('\n\n')
                   if l.strip() and not l.startswith('<!--') and not l.startswith('[')]
    # Limpar instrucoes existentes
    for idx in [41, 42, 43]:
        if idx < len(doc.paragraphs):
            doc.paragraphs[idx].clear()
    # Colocar primeiro paragrafo no slot 41
    if intro_paras:
        doc.paragraphs[41].add_run(intro_paras[0])
    # Paragrafos adicionais: inserir apos 41
    if len(intro_paras) > 1:
        ref_element = doc.paragraphs[41]._element
        for para_text in intro_paras[1:]:
            from docx.oxml.ns import qn
            new_p = ref_element.makeelement(qn('w:p'), {})
            run_el = new_p.makeelement(qn('w:r'), {})
            t_el = run_el.makeelement(qn('w:t'), {})
            t_el.text = para_text
            t_el.set(qn('xml:space'), 'preserve')
            run_el.append(t_el)
            new_p.append(run_el)
            ref_element.addnext(new_p)
            ref_element = new_p

    # === MATERIAL E METODOS (par 45 = titulo MANTER, par 47-52 = instrucoes) ===
    mm_content = sections.get('Material e Metodos', '')
    mm_blocks = re.split(r'\n\n+', mm_content)
    mm_blocks = [b.strip() for b in mm_blocks if b.strip() and not b.startswith('<!--') and not b.startswith('[')]

    # Limpar instrucoes
    for idx in [47, 48, 49, 50, 51, 52]:
        if idx < len(doc.paragraphs):
            doc.paragraphs[idx].clear()

    # Inserir M&M apos titulo (par 45)
    # Encontrar o elemento do titulo de M&M
    mm_title_el = doc.paragraphs[45]._element
    ref_el = mm_title_el

    from docx.oxml.ns import qn
    for block in mm_blocks:
        clean = block.strip()
        if not clean:
            continue
        # Detectar subtitulo
        is_subtitle = clean.startswith('**') and '**' in clean[2:]

        new_p = ref_el.makeelement(qn('w:p'), {})

        if is_subtitle:
            title_text = re.sub(r'\*\*(.+?)\*\*', r'\1', clean)
            # Adicionar recuo e negrito
            pPr = new_p.makeelement(qn('w:pPr'), {})
            ind = pPr.makeelement(qn('w:ind'), {qn('w:firstLine'): '709'})
            pPr.append(ind)
            new_p.append(pPr)

            run_el = new_p.makeelement(qn('w:r'), {})
            rPr = run_el.makeelement(qn('w:rPr'), {})
            b_el = rPr.makeelement(qn('w:b'), {})
            rPr.append(b_el)
            run_el.append(rPr)
            t_el = run_el.makeelement(qn('w:t'), {})
            t_el.text = title_text
            run_el.append(t_el)
            new_p.append(run_el)
        else:
            # Limpar markdown
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', clean)
            text = re.sub(r'\*(.+?)\*', r'\1', text)
            text = re.sub(r'`(.+?)`', r'\1', text)
            # Remover linhas de tabela
            if text.startswith('|'):
                lines = text.split('\n')
                rows = [l for l in lines if l.strip().startswith('|') and '---' not in l]
                for row in rows:
                    cells = [c.strip() for c in row.split('|') if c.strip()]
                    row_p = ref_el.makeelement(qn('w:p'), {})
                    r_el = row_p.makeelement(qn('w:r'), {})
                    t_el = r_el.makeelement(qn('w:t'), {})
                    t_el.text = ' | '.join(cells)
                    t_el.set(qn('xml:space'), 'preserve')
                    r_el.append(t_el)
                    row_p.append(r_el)
                    ref_el.addnext(row_p)
                    ref_el = row_p
                continue

            run_el = new_p.makeelement(qn('w:r'), {})
            t_el = run_el.makeelement(qn('w:t'), {})
            t_el.text = text
            t_el.set(qn('xml:space'), 'preserve')
            run_el.append(t_el)
            new_p.append(run_el)

        ref_el.addnext(new_p)
        ref_el = new_p

    # === RESULTADOS E DISCUSSAO (par 54 = titulo) ===
    rd_content = sections.get('Resultados e Discussao', '')
    rd_blocks = re.split(r'\n\n+', rd_content)
    rd_blocks = [b.strip() for b in rd_blocks if b.strip() and not b.startswith('<!--') and not b.startswith('[')]

    for idx in [56, 57]:
        if idx < len(doc.paragraphs):
            doc.paragraphs[idx].clear()

    rd_title_el = doc.paragraphs[54]._element
    ref_el = rd_title_el
    for block in rd_blocks:
        clean = block.strip()
        if not clean:
            continue
        is_subtitle = clean.startswith('**') and '**' in clean[2:]
        new_p = ref_el.makeelement(qn('w:p'), {})

        if is_subtitle:
            title_text = re.sub(r'\*\*(.+?)\*\*', r'\1', clean)
            pPr = new_p.makeelement(qn('w:pPr'), {})
            ind = pPr.makeelement(qn('w:ind'), {qn('w:firstLine'): '709'})
            pPr.append(ind)
            new_p.append(pPr)
            run_el = new_p.makeelement(qn('w:r'), {})
            rPr = run_el.makeelement(qn('w:rPr'), {})
            b_el = rPr.makeelement(qn('w:b'), {})
            rPr.append(b_el)
            run_el.append(rPr)
            t_el = run_el.makeelement(qn('w:t'), {})
            t_el.text = title_text
            run_el.append(t_el)
            new_p.append(run_el)
        elif clean.startswith('|'):
            lines = clean.split('\n')
            rows = [l for l in lines if l.strip().startswith('|') and '---' not in l]
            for row in rows:
                cells = [c.strip() for c in row.split('|') if c.strip()]
                row_p = ref_el.makeelement(qn('w:p'), {})
                r_el = row_p.makeelement(qn('w:r'), {})
                t_el = r_el.makeelement(qn('w:t'), {})
                t_el.text = ' | '.join(cells)
                t_el.set(qn('xml:space'), 'preserve')
                r_el.append(t_el)
                row_p.append(r_el)
                ref_el.addnext(row_p)
                ref_el = row_p
            continue
        else:
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', clean)
            text = re.sub(r'\*(.+?)\*', r'\1', text)
            text = re.sub(r'`(.+?)`', r'\1', text)
            run_el = new_p.makeelement(qn('w:r'), {})
            t_el = run_el.makeelement(qn('w:t'), {})
            t_el.text = text
            t_el.set(qn('xml:space'), 'preserve')
            run_el.append(t_el)
            new_p.append(run_el)

        ref_el.addnext(new_p)
        ref_el = new_p

    # === CONCLUSAO (par 59 = titulo) ===
    conc_content = sections.get('Conclusao', '')
    conc_blocks = re.split(r'\n\n+', conc_content)
    conc_blocks = [b.strip() for b in conc_blocks if b.strip() and not b.startswith('<!--') and not b.startswith('[')]

    for idx in [61, 62]:
        if idx < len(doc.paragraphs):
            doc.paragraphs[idx].clear()

    # Mudar titulo pra "Conclusão" (sem "(ões)")
    doc.paragraphs[59].clear()
    run = doc.paragraphs[59].add_run('Conclusão')
    run.bold = True

    conc_title_el = doc.paragraphs[59]._element
    ref_el = conc_title_el
    for block in conc_blocks:
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', block.strip())
        text = re.sub(r'\*(.+?)\*', r'\1', text)
        text = re.sub(r'`(.+?)`', r'\1', text)
        if not text:
            continue
        new_p = ref_el.makeelement(qn('w:p'), {})
        run_el = new_p.makeelement(qn('w:r'), {})
        t_el = run_el.makeelement(qn('w:t'), {})
        t_el.text = text
        t_el.set(qn('xml:space'), 'preserve')
        run_el.append(t_el)
        new_p.append(run_el)
        ref_el.addnext(new_p)
        ref_el = new_p

    # === REFERENCIAS (par 69 = titulo) ===
    refs_content = sections.get('Referencias', '')
    ref_entries = [l.strip() for l in refs_content.split('\n')
                   if l.strip() and not l.startswith('<!--') and not l.startswith('[') and not l.startswith('---')]

    for idx in [71, 72]:
        if idx < len(doc.paragraphs):
            doc.paragraphs[idx].clear()

    refs_title_el = doc.paragraphs[69]._element
    ref_el = refs_title_el
    for entry in ref_entries:
        clean = re.sub(r'\*\*(.+?)\*\*', r'\1', entry)
        clean = re.sub(r'\*(.+?)\*', r'\1', clean)
        if not clean:
            continue
        new_p = ref_el.makeelement(qn('w:p'), {})
        run_el = new_p.makeelement(qn('w:r'), {})
        t_el = run_el.makeelement(qn('w:t'), {})
        t_el.text = clean
        t_el.set(qn('xml:space'), 'preserve')
        run_el.append(t_el)
        new_p.append(run_el)
        ref_el.addnext(new_p)
        ref_el = new_p

    # === Salvar ===
    doc.save(str(OUT_PATH))
    print(f"Salvo em {OUT_PATH}")
    print(f"Total paragrafos: {len(doc.paragraphs)}")


if __name__ == '__main__':
    main()
